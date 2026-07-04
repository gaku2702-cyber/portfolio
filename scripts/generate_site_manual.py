# -*- coding: utf-8 -*-
"""A4・目次付き Word 解説書を生成する"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "ポートフォリオサイト運用解説書.docx"
FONT = "游ゴシック"
FONT_FALLBACK = "Yu Gothic"


def set_a4(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        hs.font.color.rgb = RGBColor(0x1B, 0x32, 0x64)
        if level == 1:
            hs.font.size = Pt(16)
            hs.paragraph_format.space_before = Pt(18)
            hs.paragraph_format.space_after = Pt(10)
        elif level == 2:
            hs.font.size = Pt(13)
            hs.paragraph_format.space_before = Pt(14)
            hs.paragraph_format.space_after = Pt(8)
        else:
            hs.font.size = Pt(11)
            hs.paragraph_format.space_before = Pt(10)
            hs.paragraph_format.space_after = Pt(6)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    hint = OxmlElement("w:t")
    hint.text = "（目次を更新するには、この行を右クリック →「フィールドの更新」）"
    run._r.append(hint)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.bold = bold
    return para


def bullet(doc, text):
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(text, style="List Number")
    for run in para.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return para


def code_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    return para


def table_simple(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = FONT
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.add_paragraph()
    return table


def build() -> None:
    doc = Document()
    set_a4(doc)
    set_default_font(doc)

    # ===== 表紙 =====
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ポートフォリオサイト\n運用・更新 解説書")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.name = FONT
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    tr.font.color.rgb = RGBColor(0x1B, 0x32, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("DTP × DX Engineer Portfolio\nサイトの見方・編集方法・公開手順")
    sr.font.size = Pt(12)
    sr.font.name = FONT
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("対象サイト：https://gaku2702-cyber.github.io/portfolio/\n"
                      "プロジェクトフォルダ：D:\\開発アプリ\\ポートフォリオ計画\n"
                      "用紙：A4　／　印刷・配布用")
    mr.font.size = Pt(10)
    mr.font.name = FONT
    mr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    doc.add_page_break()

    # ===== 目次 =====
    doc.add_heading("目次", level=1)
    p(doc, "Word で開いたあと、目次が空の場合は次の操作を行ってください。")
    numbered(doc, "目次のフィールド部分を右クリックする")
    numbered(doc, "「フィールドの更新」を選ぶ")
    numbered(doc, "「目次全体を更新する」を選んで OK")
    doc.add_paragraph()
    toc_para = doc.add_paragraph()
    add_toc_field(toc_para)
    doc.add_page_break()

    # ===== 第1章 =====
    doc.add_heading("第1章　このサイトの仕組み", level=1)
    p(doc, "このポートフォリオサイトは、WordPress のような管理画面はありません。"
          "HTML という形式のファイルに文章・画像・リンクが直接書かれており、"
          "そのファイルを編集して GitHub にアップロード（push）すると、"
          "インターネット上の公開サイトが更新されます。")

    doc.add_heading("1-1　管理画面はない", level=2)
    p(doc, "一般的なホームページ作成サービスと異なり、ブラウザからログインして"
          "ボタンで編集する画面はありません。代わりに、パソコン上のフォルダ内の"
          "ファイルを Cursor やメモ帳などで開いて編集します。")

    doc.add_heading("1-2　公開の流れ（3ステップ）", level=2)
    numbered(doc, "ローカル（自分の PC）で index.html などを編集する")
    numbered(doc, "ブラウザで index.html を開き、表示を確認する")
    numbered(doc, "git push で GitHub に送る → 1〜2 分で公開サイトに反映")

    doc.add_heading("1-3　公開 URL", level=2)
    p(doc, "https://gaku2702-cyber.github.io/portfolio/")

    # ===== 第2章 =====
    doc.add_heading("第2章　フォルダとファイルの地図", level=1)
    p(doc, "プロジェクトの本体は次のフォルダにあります。")
    code_block(doc, r"D:\開発アプリ\ポートフォリオ計画")

    table_simple(
        doc,
        ["ファイル／フォルダ", "役割"],
        [
            ["index.html", "メインのポートフォリオページ（ほとんどの編集はここ）"],
            ["imposition-demo.html", "JFX200 面付シミュレーターの体験デモ"],
            ["assets/docs/pman-google-workspace-dx-roadmap.pdf", "P-MAN × Google Workspace 提案書 PDF（表紙クリックで直接表示）"],
            ["assets/covers/", "開発実績セクションの表紙画像"],
            ["assets/docs/", "提案書 PDF 本体"],
            ["Project03_… など", "参考資料・別プロジェクト（サイト本体ではない）"],
        ],
    )

    p(doc, "※ サイトの見た目を変えたいときは、まず index.html を疑ってください。", bold=True)

    # ===== 第3章 =====
    doc.add_heading("第3章　index.html の構成（どこに何があるか）", level=1)
    p(doc, "index.html は1つの長いファイルです。"
          "次の見出しコメント（<!-- ... -->）を検索すると、該当セクションにすぐ移動できます。")

    table_simple(
        doc,
        ["検索キーワード", "セクション", "内容"],
        [
            ["HERO", "トップ画面", "キャッチコピー・メインビジュアル"],
            ["CONCEPT", "コンセプト", "強み・特徴の説明"],
            ["SERVICES & WORKS", "開発実績", "表紙ギャラリー（Works）"],
            ["CONTACT", "お問い合わせ", "フォーム"],
            ["TOOL DETAIL MODAL", "詳細ポップアップ", "Word 変換ツールの説明"],
        ],
    )

    p(doc, "Cursor で index.html を開き、Ctrl + F（検索）で上記の英語を入力すると、"
          "該当行にジャンプできます。")

    # ===== 第4章 =====
    doc.add_heading("第4章　開発実績（Works）の変更方法", level=1)
    p(doc, "開発実績は「表紙画像を並べる」形式になっています。"
          "画像ファイルと index.html の2か所を理解すれば、ほとんどの更新ができます。")

    doc.add_heading("4-1　表紙画像だけ差し替える（最も簡単）", level=2)
    p(doc, "アプリ画面や PDF 表紙のスクリーンショットを、次のフォルダに同じファイル名で上書き保存します。")
    code_block(doc, "assets/covers/")

    table_simple(
        doc,
        ["ファイル名", "対象"],
        [
            ["imposition-demo.jpg", "JFX200 面付シミュレーター"],
            ["word-tag-extractor.jpg", "Word タグ付きテキスト化ツール"],
            ["indesign-qa.jpg", "入稿チェック自動化スクリプト"],
            ["pman-dx-proposal.jpg", "P-MAN × Google Workspace 連結 DX 提案書"],
            ["dx-consulting.jpg", "DX 化ヒアリング & ロードマップ策定"],
        ],
    )

    p(doc, "重要：ファイル名を変えないでください。"
          "index.html はファイル名で画像を参照しているため、"
          "名前を変えると index.html 側も書き換える必要があります。", bold=True)

    doc.add_heading("4-2　タイトル・リンク・ボタン文言を変える", level=2)
    p(doc, "index.html 内の「SERVICES & WORKS」セクション（449 行目付近）を編集します。"
          "各実績は <!-- Work 1 --> 〜 <!-- Work 5 --> のブロックに分かれています。")

    table_simple(
        doc,
        ["変更したい内容", "編集する HTML の部分"],
        [
            ["クリックしたときのリンク先", 'href="..." の URL'],
            ["表紙画像", 'src="assets/covers/..."'],
            ["マウスを乗せたときの文言", "体験デモを開く などの日本語"],
            ["カテゴリ名（英語）", "DTP Automation など"],
            ["作品名", "<h3> タグ内の日本語"],
        ],
    )

    doc.add_heading("4-3　実績を1件追加する", level=2)
    numbered(doc, "表紙画像を assets/covers/ に保存（例：new-project.jpg）")
    numbered(doc, "index.html の Cover gallery 内で、既存の Work 1 ブロックをコピー")
    numbered(doc, "画像パス・タイトル・href を新しい内容に書き換え")
    numbered(doc, "ブラウザで表示確認 → git push")

    doc.add_heading("4-4　Word 変換ツールの詳細ポップアップ", level=2)
    p(doc, "「詳細・処理例を見る」を押したときに開く説明は、"
          "index.html の「TOOL DETAIL MODAL」（935 行目付近）にあります。"
          "機能リストや Before/After の例文はここを編集してください。")

    # ===== 第5章 =====
    doc.add_heading("第5章　その他のページ・ファイル", level=1)

    table_simple(
        doc,
        ["編集したい内容", "ファイル"],
        [
            ["面付デモの計算・UI", "imposition-demo.html"],
            ["提案書 PDF の中身", "assets/docs/ 内の PDF を差し替え（表紙は assets/covers/pman-dx-proposal.jpg）"],
        ],
    )

    p(doc, "提案書 PDF を差し替えた場合、表紙画像も合わせて更新すると、"
          "開発実績のサムネイルと中身が一致します。")

    # ===== 第6章 =====
    doc.add_heading("第6章　変更内容の確認方法（ローカルプレビュー）", level=1)
    numbered(doc, "エクスプローラーで D:\\開発アプリ\\ポートフォリオ計画 を開く")
    numbered(doc, "index.html をダブルクリック（Chrome または Edge で開く）")
    numbered(doc, "表示を確認する")
    numbered(doc, "古い表示のままなら F5（更新）または Ctrl + F5（強制更新）")

    p(doc, "画像を差し替えた直後は、ブラウザのキャッシュで古い画像が残ることがあります。"
          "Ctrl + F5 でほとんど解消します。")

    # ===== 第7章 =====
    doc.add_heading("第7章　インターネット上のサイトを更新する（git push）", level=1)
    p(doc, "ローカルで問題なければ、GitHub に送って公開サイトを更新します。"
          "Cursor のターミナル（PowerShell）で次を実行します。")

    code_block(doc, 'cd "D:\\開発アプリ\\ポートフォリオ計画"')
    code_block(doc, "git add .")
    code_block(doc, 'git commit -m "開発実績の表紙を更新"')
    code_block(doc, "git push")

    p(doc, "push 後 1〜2 分待ってから、公開 URL をブラウザで開いて確認してください。")

    doc.add_heading("7-1　git コマンドの意味（参考）", level=2)
    table_simple(
        doc,
        ["コマンド", "意味"],
        [
            ["git add .", "変更したファイルを「コミット候補」に載せる"],
            ["git commit -m \"…\"", "変更内容にメモを付けて記録する"],
            ["git push", "GitHub に送り、公開サイトを更新する"],
        ],
    )

    # ===== 第8章 =====
    doc.add_heading("第8章　よくあるトラブルと対処", level=1)

    doc.add_heading("8-1　画像を変えたのにサイトが変わらない", level=2)
    bullet(doc, "assets/covers/ のファイル名が index.html の src= と一致しているか確認")
    bullet(doc, "git push まで実行したか確認")
    bullet(doc, "ブラウザで Ctrl + F5（強制更新）を試す")

    doc.add_heading("8-2　リンクを押してもページが開かない", level=2)
    bullet(doc, "href= のファイル名のスペルミスがないか確認")
    bullet(doc, "リンク先ファイル（例：imposition-demo.html）が同じフォルダにあるか確認")

    doc.add_heading("8-3　HTML が難しくて不安", level=2)
    p(doc, "次の2つだけ覚えれば、日常の更新は十分です。")
    bullet(doc, "表紙だけ変える → assets/covers/ に同じ名前で上書き")
    bullet(doc, "タイトルだけ変える → index.html の <h3> の日本語部分を書き換え")

    # ===== 付録 =====
    doc.add_heading("付録A　表紙画像の推奨サイズ", level=1)
    p(doc, "比率 16:9（横長）を推奨します。例：960 × 540 px、1920 × 1080 px。"
          "形式は JPG または PNG。ファイルサイズは 500 KB 以下程度が読み込みに有利です。")

    doc.add_heading("付録B　用語ミニ辞典", level=1)
    table_simple(
        doc,
        ["用語", "意味"],
        [
            ["HTML", "Web ページの設計図となるファイル形式"],
            ["index.html", "サイトのトップページ本体"],
            ["push", "変更を GitHub に送って公開する操作"],
            ["ローカル", "自分の PC 上（インターネット公開前）"],
            ["表紙ギャラリー", "開発実績を表紙画像で並べるレイアウト"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 以上 —")
    fr.font.name = FONT
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    fr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build()
